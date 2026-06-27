"""Knowledge base upload routes."""

import os
from io import BytesIO
from typing import Any

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.security import get_current_tenant_id, require_role
from app.db.models import KnowledgeBaseDoc
from app.db.session import get_db
from app.modules.rag_knowledge_base.service import add_chunks, add_document

router = APIRouter(prefix="/api/v1/admin", tags=["knowledge_base"])

MAX_KB_DOCS = 10
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
MAX_PDF_PAGES = 500
ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _get_safe_extension(filename: str) -> str:
    """Extract safe file extension without path traversal risk."""
    base = os.path.basename(filename)
    # Remove any null bytes
    base = base.replace("\x00", "")
    if "." not in base:
        return ""
    # Use only the last extension
    return base.rsplit(".", 1)[-1].lower().strip()


def _extract_text(filename: str, content: bytes) -> str:
    """Extract text from PDF, TXT, or DOCX."""
    ext = _get_safe_extension(filename)
    if ext == "txt":
        return content.decode("utf-8", errors="ignore")
    if ext == "pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(BytesIO(content))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError(f"PDF exceeds maximum page limit ({MAX_PDF_PAGES})")
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n".join(parts)
    if ext == "docx":
        from docx import Document

        doc = Document(BytesIO(content))
        parts = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style = p.style
            style_name = (style.name or "").lower() if style else ""
            style_id = (style.style_id or "").lower() if style else ""
            # Recognise Word heading styles in any common locale
            is_heading = (
                style_name.startswith("heading")
                or style_name.startswith("заголовок")
                or style_name.startswith("titre")
                or style_name.startswith("überschrift")
                or style_id.startswith("heading")
            )
            if is_heading:
                # Prefer the numeric level from style name/id
                level_str = "".join(ch for ch in (style.name or "") if ch.isdigit())
                if not level_str:
                    level_str = "".join(ch for ch in (style.style_id or "") if ch.isdigit())
                level = max(1, min(6, int(level_str) if level_str else 1))
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        return "\n\n".join(parts)
    raise ValueError(f"Unsupported file type: .{ext}")


def _split_by_sentences(text: str, max_chunk_size: int, overlap: int) -> list[str]:
    """Split a long paragraph into sentence-bounded chunks with overlap."""
    import re

    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    chunks: list[str] = []
    current = ""
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        candidate = f"{current} {sentence}".strip() if current else sentence
        if len(candidate) <= max_chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current)
            if overlap and current:
                tail = current[-overlap:]
                # Try to start overlap from a sentence boundary
                boundary = tail.find(". ")
                if boundary != -1:
                    current = tail[boundary + 2:] + " " + sentence
                else:
                    current = tail + " " + sentence
                current = current.strip()
            else:
                current = sentence
    if current:
        chunks.append(current)
    return chunks


def _smart_chunk(text: str, max_chunk_size: int = 700, overlap: int = 100) -> list[str]:
    """Semantic chunking preserving section boundaries, FAQ Q&A pairs and scenarios.

    Strategy:
    1. Group text by major section headers (1., 3.1., 4., etc. — mostly UPPERCASE titles)
    2. Inside each group — extract FAQ pairs (❓ question + following answer lines)
    3. Extract scenario blocks ("Сценарий N:" + following dialogue)
    4. Keep plain content grouped by paragraph breaks; never merge across sections
    5. Split oversized plain text by paragraphs (not arbitrary sentences)
    """
    import re

    text = re.sub(r"\n{3,}", "\n\n", text.strip())
    lines = text.split("\n")

    def _is_major_header(line: str) -> bool:
        line = line.strip()
        if not line:
            return False
        if line.startswith("#"):
            return True
        m = re.match(r"^(\d+(\.\d+)*)\.?\s+(.+)$", line)
        if m:
            title = m.group(3)
            letters = [c for c in title if c.isalpha()]
            if letters:
                upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
                if upper_ratio > 0.7:
                    return True
            if len(line) < 40:
                return True
        return False

    def _is_faq(line: str) -> bool:
        line = line.strip()
        return line.startswith("❓") or line.startswith("Вопрос:") or line.startswith("? ")

    def _is_scenario(line: str) -> bool:
        line = line.strip()
        return line.startswith("Сценарий ") and ":" in line

    # Phase 1: group by major headers
    groups: list[list[str]] = []
    current_group: list[str] = []
    for line in lines:
        stripped = line.strip()
        if _is_major_header(stripped):
            if current_group:
                groups.append(current_group)
            current_group = [line]
        else:
            current_group.append(line)
    if current_group:
        groups.append(current_group)

    # Phase 2: process each group
    chunks: list[str] = []
    pending_empty_header: str | None = None

    for group in groups:
        header = group[0].strip()
        body = group[1:]

        if not body or not any(b.strip() for b in body):
            # Empty group — merge header with next group
            if pending_empty_header:
                pending_empty_header += " | " + header
            else:
                pending_empty_header = header
            continue

        if pending_empty_header:
            header = pending_empty_header + " | " + header
            pending_empty_header = None

        # Rebuild body text and split into sub-blocks
        sub_blocks: list[tuple[str, str]] = []
        current_sub: list[str] = []
        for line in body:
            stripped = line.strip()
            if _is_faq(stripped) or _is_scenario(stripped):
                if current_sub:
                    sub_blocks.append(("\n".join(current_sub).strip(), "content"))
                    current_sub = []
                sub_blocks.append((line, "faq" if _is_faq(stripped) else "scenario"))
            else:
                current_sub.append(line)
        if current_sub:
            sub_blocks.append(("\n".join(current_sub).strip(), "content"))

        # Build FAQ pairs and scenario blocks
        faq_pairs: list[str] = []
        i = 0
        while i < len(sub_blocks):
            block, btype = sub_blocks[i]
            if btype == "faq":
                answer_parts: list[str] = []
                j = i + 1
                while j < len(sub_blocks) and sub_blocks[j][1] == "content":
                    answer_parts.append(sub_blocks[j][0])
                    j += 1
                answer = "\n".join(answer_parts).strip()
                faq_pairs.append(f"{block}\n{answer}" if answer else block)
                i = j
            elif btype == "scenario":
                scenario_parts = [block]
                j = i + 1
                while j < len(sub_blocks) and sub_blocks[j][1] == "content":
                    scenario_parts.append(sub_blocks[j][0])
                    j += 1
                chunks.append("\n".join(scenario_parts).strip())
                i = j
            else:
                i += 1

        # Collect plain content blocks not absorbed by FAQ/scenario
        plain_blocks: list[str] = []
        i = 0
        while i < len(sub_blocks):
            block, btype = sub_blocks[i]
            if btype == "content":
                plain_blocks.append(block)
                i += 1
            elif btype == "faq":
                j = i + 1
                while j < len(sub_blocks) and sub_blocks[j][1] == "content":
                    j += 1
                i = j
            elif btype == "scenario":
                j = i + 1
                while j < len(sub_blocks) and sub_blocks[j][1] == "content":
                    j += 1
                i = j

        plain_text = "\n\n".join(b for b in plain_blocks if b).strip()
        if plain_text:
            paragraphs = [p.strip() for p in plain_text.split("\n\n") if p.strip()]
            current = header
            last_para = ""
            for para in paragraphs:
                candidate = (
                    f"{current}\n\n{para}".strip()
                    if current != header
                    else f"{header}\n\n{para}".strip()
                )
                # If a single paragraph is oversized, split it by sentences.
                if len(para) > max_chunk_size:
                    if current and current != header:
                        chunks.append(current)
                        current = header
                    for sentence_chunk in _split_by_sentences(para, max_chunk_size, overlap):
                        chunks.append(f"{header}\n\n{sentence_chunk}".strip())
                    last_para = para
                    continue
                if len(candidate) <= max_chunk_size:
                    current = candidate
                    last_para = para
                else:
                    if current and current != header:
                        chunks.append(current)
                    if overlap and last_para:
                        current = f"{header}\n\n{last_para}\n\n{para}".strip()
                    else:
                        current = f"{header}\n\n{para}".strip()
                    last_para = para
            if current and current != header:
                chunks.append(current)

        for pair in faq_pairs:
            chunks.append(pair)

    if pending_empty_header:
        chunks.append(pending_empty_header)

    # Clean up: deduplicate while preserving order
    seen: set[str] = set()
    result: list[str] = []
    for c in chunks:
        c = c.strip()
        if c and c not in seen:
            seen.add(c)
            result.append(c)
    return result


@router.post("/knowledge")
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    tenant_id: int = Depends(get_current_tenant_id),
    user: dict[str, Any] = Depends(require_role("tenant_admin", "superadmin")),
) -> dict[str, Any]:
    """Upload a PDF/TXT/DOCX document."""
    # Validate filename and extension
    ext = _get_safe_extension(file.filename or "")
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Invalid file type. Allowed: PDF, TXT, DOCX")

    # Validate MIME type if provided
    if file.content_type and file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(status_code=400, detail="Invalid content type")

    # SECURITY: Read file with streaming and early size check
    # to prevent memory exhaustion from huge uploads
    content = bytearray()
    chunk_size = 64 * 1024  # 64 KB chunks
    total_read = 0
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        total_read += len(chunk)
        if total_read > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail=f"File too large (max {MAX_FILE_SIZE // (1024 * 1024)}MB)")
        content.extend(chunk)

    # Check document count limit
    doc_count = await db.scalar(
        select(func.count(KnowledgeBaseDoc.id)).where(KnowledgeBaseDoc.tenant_id == tenant_id)
    ) or 0
    if doc_count >= MAX_KB_DOCS:
        return {"error": f"Document limit reached (max {MAX_KB_DOCS} files). Delete old files first."}

    text = _extract_text(file.filename, content)
    if not text.strip():
        return {"error": "Could not extract text from file"}

    doc = await add_document(db, tenant_id, file.filename)
    # Smart chunking: split by logical blocks (FAQ questions, sections), then merge small ones
    chunks = _smart_chunk(text, max_chunk_size=1200, overlap=150)

    await add_chunks(db, tenant_id, doc.id, chunks)
    return {"doc_id": doc.id, "chunks": len(chunks)}


@router.get("/knowledge")
async def list_documents(
    db: AsyncSession = Depends(get_db),
    tenant_id: int = Depends(get_current_tenant_id),
    user: dict[str, Any] = Depends(require_role("tenant_admin", "superadmin")),
) -> list[dict[str, Any]]:
    """List uploaded knowledge base documents."""
    result = await db.execute(
        select(KnowledgeBaseDoc).where(KnowledgeBaseDoc.tenant_id == tenant_id)
    )
    docs = result.scalars().all()
    return [
        {"id": d.id, "filename": d.filename, "status": d.status, "created_at": d.created_at.isoformat()}
        for d in docs
    ]


@router.get("/knowledge/search")
async def search_knowledge_debug(
    q: str,
    db: AsyncSession = Depends(get_db),
    tenant_id: int = Depends(get_current_tenant_id),
    user: dict[str, Any] = Depends(require_role("tenant_admin", "superadmin")),
) -> dict[str, Any]:
    """Debug RAG search — returns raw chunks with scores."""
    from app.modules.rag_knowledge_base.service import search_knowledge_with_scores
    results = await search_knowledge_with_scores(db, tenant_id, q, top_k=5)
    return {
        "query": q,
        "chunks_found": len(results),
        "results": [
            {"content": r["content"][:300], "distance": r["distance"]} for r in results
        ],
    }


@router.get("/knowledge/{doc_id}/chunks")
async def list_document_chunks(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    tenant_id: int = Depends(get_current_tenant_id),
    user: dict[str, Any] = Depends(require_role("tenant_admin", "superadmin")),
) -> list[dict[str, Any]]:
    """List all chunks for a document with sizes."""
    from app.db.models import KnowledgeBaseChunk
    result = await db.execute(
        select(KnowledgeBaseChunk)
        .where(
            KnowledgeBaseChunk.doc_id == doc_id,
            KnowledgeBaseChunk.tenant_id == tenant_id,
        )
        .order_by(KnowledgeBaseChunk.id)
    )
    chunks = result.scalars().all()
    return [
        {
            "id": c.id,
            "size": len(c.content),
            "has_embedding": c.embedding is not None,
            "preview": c.content[:200],
        }
        for c in chunks
    ]


@router.delete("/knowledge/{doc_id}")
async def delete_document(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    tenant_id: int = Depends(get_current_tenant_id),
    user: dict[str, Any] = Depends(require_role("tenant_admin", "superadmin")),
) -> dict[str, str]:
    """Delete a knowledge base document and all its chunks."""
    from sqlalchemy import delete
    from app.db.models import KnowledgeBaseChunk, KnowledgeBaseDoc

    doc = await db.scalar(
        select(KnowledgeBaseDoc).where(
            KnowledgeBaseDoc.id == doc_id,
            KnowledgeBaseDoc.tenant_id == tenant_id,
        )
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    await db.execute(
        delete(KnowledgeBaseChunk).where(
            KnowledgeBaseChunk.doc_id == doc_id,
            KnowledgeBaseChunk.tenant_id == tenant_id,
        )
    )
    await db.delete(doc)
    await db.commit()
    return {"status": "deleted"}
